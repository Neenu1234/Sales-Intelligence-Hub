from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin   = Inches(1.2)
section.right_margin  = Inches(1.2)

# ── Helper functions ──────────────────────────────────────────────────────────
def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb) if level == 1 else RGBColor(0x1e, 0x40, 0xaf)
    return p

def body(text, bold_parts=None):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.space_before = Pt(2)
    return p

def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    return p

def italic_cite(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_after  = Pt(8)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data rows
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            cells[ci].paragraphs[0].paragraph_format.space_after = Pt(2)
    doc.add_paragraph()

def divider():
    doc.add_paragraph("─" * 72)

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Interactive Online Sales Performance &\nCustomer Behavior Dashboard")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub.add_run("Milestone 1: Exploratory Data Analysis & Prototype Layout")
run2.font.size = Pt(13)
run2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = meta.add_run("Data Visualization  |  April 2026")
run3.font.size = Pt(11)
run3.italic = True

doc.add_paragraph()
divider()
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION & INDUSTRY MOTIVATION
# ══════════════════════════════════════════════════════════════════════════════
heading("1. Introduction & Industry Motivation")

heading("1.1 Project Overview", level=2)

body(
    "This project builds an interactive data visualization dashboard for online sales "
    "performance and customer behavior analysis. Using a real-world e-commerce transactions "
    "dataset with nearly 50,000 order records spanning multiple countries, product categories, "
    "payment methods, and sales channels, the dashboard enables business stakeholders to "
    "explore revenue trends, identify top-performing markets, understand customer loyalty "
    "patterns, and uncover operational insights — all through interactive visual charts "
    "rather than raw spreadsheets."
)

body(
    "The dashboard is implemented in Streamlit and includes eleven coordinated visualizations "
    "covering time-series trends, geographic maps, product category breakdowns, customer "
    "segmentation, payment flow analysis, and an AI-powered natural language assistant. "
    "Every chart responds dynamically to sidebar filters, allowing users to drill into "
    "any combination of country, category, sales channel, or time period in real time."
)

body(
    "The significance of this project lies in a specific and underserved gap: most retail "
    "businesses already have data analysts and existing dashboards — but those dashboards "
    "are built by analysts for executives, showing only what the analyst decided to show. "
    "If a manager wants to ask a different question, they have to go back to the analyst "
    "and wait. This project eliminates that bottleneck by making the dashboard fully "
    "self-service — any stakeholder can filter, explore, and even ask questions in plain "
    "English through the built-in AI assistant, without involving a data team at all."
)

heading("1.2 Why Online Sales Data? Why Visualization?", level=2)

body(
    "Online sales is one of the richest and most commercially important domains for data "
    "visualization, for three specific reasons."
)

p = doc.add_paragraph()
run = p.add_run("1. The data is inherently multi-dimensional. ")
run.bold = True
p.add_run(
    "Every transaction in an online retail dataset carries information across time "
    "(when was it placed?), geography (which country?), product (which category?), "
    "behavior (which payment method, which channel?), and customer identity (who bought it, "
    "have they bought before?). No single number or table can communicate all of these "
    "dimensions at once. Visualization is the only way to let a person explore all of "
    "them simultaneously and spot the relationships between them."
)

p = doc.add_paragraph()
run = p.add_run("2. The stakes of misreading this data are high. ")
run.bold = True
p.add_run(
    "The global e-commerce market exceeded $5.8 trillion in 2023 (Statista, 2024). "
    "At that scale, a manager who cannot quickly see that one market is declining, or "
    "that a product category is underperforming, or that loyal customers are churning, "
    "is making decisions blind. Incorrect or delayed decisions translate directly into "
    "lost revenue and wasted marketing spend. Visualization is not decorative here — "
    "it is the mechanism through which the right decision gets made at the right time."
)

p = doc.add_paragraph()
run = p.add_run("3. The audience is non-technical. ")
run.bold = True
p.add_run(
    "The people who need these insights most — store managers, marketing leads, product "
    "owners — are not data engineers. They cannot run SQL queries or read a pandas "
    "DataFrame. Data visualization bridges that gap by encoding the data visually in a "
    "form that requires no technical knowledge to interpret. A choropleth map communicates "
    "geographic revenue concentration instantly. A sunburst chart shows payment method "
    "and channel breakdown in one glance. An interactive filter lets anyone ask their "
    "own question without writing a line of code."
)

heading("1.3 The Real Problem This Solves", level=2)

body(
    "Imagine you are a sales manager at an online retail company. It is Monday morning "
    "and your CEO wants to know: why did revenue drop last month, which markets are "
    "underperforming, and which customers are at risk of leaving? You have the data — "
    "tens of thousands of transaction records — but getting a clear answer means going "
    "to your data analyst, waiting for them to run queries, build a report, and send it "
    "back. By the time you have an answer, the moment to act has already passed."
)

body(
    "Most retail teams are not short on data. They are not even short on analysts. "
    "They are short on speed — the speed at which a business question becomes a "
    "visual answer that anyone can act on. The bottleneck is not data collection; "
    "it is the dependency between non-technical decision-makers and technical teams."
)

body(
    "This dashboard breaks that dependency. A store manager, a marketing lead, or a "
    "product owner can open it, apply their own filters, and get answers immediately — "
    "without writing a single query. And if they want to ask something that no chart "
    "covers, the AI assistant answers in plain English using the actual dataset. "
    "That is the core innovation: not just visualizing data, but making data exploration "
    "fully self-service for people who are not data professionals."
)

heading("1.4 What Makes This Dashboard Different — The AI Assistant", level=2)

body(
    "Most dashboards stop at charts. This one goes further with a built-in AI assistant "
    "that lets users ask questions about the data in plain English and get answers "
    "instantly. This is not a generic chatbot — it is specifically trained on the dataset."
)

p = doc.add_paragraph()
run = p.add_run("How it works: ")
run.bold = True
p.add_run(
    "Before any question is asked, the application automatically computes key metrics "
    "from the dataset — top revenue country, category breakdowns, return rates, payment "
    "method totals, and more — and packages all of that as a data context. When a user "
    "types a question, that context is sent to the Claude AI model (via the Anthropic API) "
    "alongside the question. Claude reads the actual computed numbers and answers based "
    "on the real data, not general knowledge. The full conversation history is maintained "
    "in session state so follow-up questions work naturally."
)

body(
    "For example, a user can type 'Which category has the highest return rate?' or "
    "'How does online revenue compare to in-store?' and receive a precise, data-backed "
    "answer in seconds — without touching a single filter or reading a single chart. "
    "This directly addresses the project's core motivation: reducing the gap between "
    "a business question and a data-driven answer."
)

heading("1.5 Research Evidence Supporting This Approach", level=2)

body("Research Paper 1 — Dashboard Value for Decision-Making")
body(
    "Sarikaya et al. (2019) studied how dashboards are used across industry and found "
    "that their primary value is enabling rapid situation awareness for people who are "
    "not data specialists. Dashboards that combine summary KPIs with interactive "
    "drill-down views significantly reduce the time from question to answer compared "
    "to traditional reporting. This directly shaped our design: the KPI row gives an "
    "immediate overview, while the charts below allow deeper exploration."
)
italic_cite(
    "Sarikaya, A., Correll, M., Bartram, L., Tory, M., & Fisher, D. (2019). What do we talk "
    "about when we talk about dashboards? IEEE Transactions on Visualization and Computer "
    "Graphics, 25(1), 682–692. https://doi.org/10.1109/TVCG.2018.2864903"
)

body("Research Paper 2 — Interactive Visualization for Analytical Insight")
body(
    "Heer & Shneiderman (2012) identified three interaction types that are now foundational "
    "in visualization design: dynamic filtering, brushing and linking, and details-on-demand "
    "via hover tooltips. They found these mechanisms help users discover patterns they would "
    "completely miss in a static table. Our dashboard implements all three across every chart."
)
italic_cite(
    "Heer, J., & Shneiderman, B. (2012). Interactive dynamics for visual analysis. Queue, "
    "10(2), 30–55. https://doi.org/10.1145/2133416.2146416"
)

body("Research Paper 3 — Customer Segmentation in Retail Analytics")
body(
    "Ngai, Xiu & Chau (2009) reviewed over 900 papers on data mining in retail CRM and "
    "found that customer segmentation based on purchase behavior consistently delivers "
    "the highest business impact. RFM analysis — grouping customers by Recency, Frequency, "
    "and Monetary value — emerged as the industry standard for retention decisions. This "
    "directly validates the customer loyalty section of our dashboard."
)
italic_cite(
    "Ngai, E.W.T., Xiu, L., & Chau, D.C.K. (2009). Application of data mining techniques in "
    "customer relationship management: A literature review and classification. Expert Systems "
    "with Applications, 36(2), 2592–2602. https://doi.org/10.1016/j.eswa.2007.10.017"
)

body("Research Paper 4 — The Business Cost of Ignoring Customer Data")
body(
    "Chen & Popovich (2003) found that companies integrating CRM analytics into daily "
    "operations consistently outperform those that do not — particularly in retention. "
    "This matters because acquiring a new customer costs 5 to 7 times more than keeping "
    "an existing one. Our customer loyalty visualization is designed to surface exactly "
    "that risk before it becomes a loss."
)
italic_cite(
    "Chen, I.J., & Popovich, K. (2003). Understanding customer relationship management (CRM): "
    "People, process and technology. Business Process Management Journal, 9(5), 672–688. "
    "https://doi.org/10.1108/14637150310496758"
)

heading("1.3 Who Will Use This Dashboard", level=2)

body("The dashboard is designed to serve multiple stakeholder roles in a retail organization:")

add_table(
    ["Stakeholder", "Primary Use Case", "Key Visualizations"],
    [
        ["Sales Manager", "Track monthly revenue trends, identify slow periods", "Revenue timeline, Waterfall chart"],
        ["Marketing Director", "Find top-performing countries and channels", "Choropleth map, Sankey flow"],
        ["Product Manager", "Understand which categories drive orders", "Treemap, Category bar chart"],
        ["CRM / Retention Team", "Identify loyal vs. at-risk customer segments", "Loyalty bubble, RFM analysis"],
        ["Operations", "Plan staffing by peak hours and days", "24-hour clock, Day-of-week heatmap"],
    ]
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 2. PROBLEM STATEMENT
# ══════════════════════════════════════════════════════════════════════════════
heading("2. Problem Statement")

heading("2.1 Core Question", level=2)

p = doc.add_paragraph()
run = p.add_run(
    "How can a retail business efficiently identify revenue trends, top-performing segments, "
    "and customer loyalty patterns in online sales data — without relying on manual SQL queries "
    "or spreadsheet exports?"
)
run.bold = True
run.font.size = Pt(12)
p.paragraph_format.space_after = Pt(8)

body("More specifically, the dashboard will answer:")
bullet("How is revenue evolving month by month, and are there seasonal peaks?")
bullet("Which product categories and countries generate the most revenue?")
bullet("Which payment methods and sales channels dominate transactions?")
bullet("How do customers segment by loyalty (repeat buyers vs. one-time buyers)?")
bullet("At what hours and days of the week do sales concentrate?")

heading("2.2 Objectives and Expected Insights", level=2)

add_table(
    ["Objective", "Visualization", "Expected Business Insight"],
    [
        ["Trend tracking", "Line/area chart by month", "Detect growth periods, seasonality, campaign effects"],
        ["Category analysis", "Treemap + bar chart", "Identify top and underperforming product categories"],
        ["Geographic performance", "Choropleth world map", "Reveal highest-value markets and expansion opportunities"],
        ["Customer loyalty", "Bubble chart + RFM", "Distinguish Champions vs. At-Risk customers"],
        ["Operations planning", "Polar clock + heatmap", "Pinpoint peak hours and days for staffing"],
        ["AI exploration", "Natural language chatbot", "Reduce friction for non-technical users"],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
# 3. DATASET DESCRIPTION
# ══════════════════════════════════════════════════════════════════════════════
heading("3. Dataset Description")

heading("3.1 Dataset Source", level=2)

body(
    "This project uses the Online Sales Dataset — a comprehensive e-commerce transactions "
    "dataset containing invoice-level order records structured to reflect real-world online "
    "retail operations."
)

add_table(
    ["Attribute", "Detail"],
    [
        ["File name", "online_sales_dataset.csv"],
        ["Format", "CSV, comma-delimited"],
        ["Total records", "~49,782 transaction rows"],
        ["Total columns", "17 attributes"],
        ["Time span", "Multiple years of transaction history"],
    ]
)

heading("3.2 Why This Dataset?", level=2)

body("This dataset was selected for three key reasons:")

p = doc.add_paragraph()
run = p.add_run("1. Multi-dimensional richness — ")
run.bold = True
p.add_run(
    "The dataset spans temporal (date, hour), geographic (country), behavioral (payment method, "
    "sales channel, return status), and product (category) dimensions. This allows construction "
    "of a wide variety of chart types — time series, choropleth maps, sunburst charts, heatmaps, "
    "and flow diagrams — all from a single source, which is essential for a comprehensive "
    "visualization project."
)

p = doc.add_paragraph()
run = p.add_run("2. Real-world industry relevance — ")
run.bold = True
p.add_run(
    "The columns directly mirror what companies like Amazon, Shopify, and Walmart collect at "
    "every transaction (order ID, invoice date, product category, quantity, unit price, "
    "customer ID, country, payment method, shipping cost, discount). Any insight produced by "
    "this dashboard could be applied directly by a retail analytics team."
)

p = doc.add_paragraph()
run = p.add_run("3. Customer-level analysis capability — ")
run.bold = True
p.add_run(
    "The presence of a CustomerID column enables RFM (Recency, Frequency, Monetary) customer "
    "segmentation — one of the most impactful analytics tasks in e-commerce, as documented by "
    "Ngai et al. (2009). With this field, we can distinguish Champions (high-frequency, "
    "high-value buyers) from At-Risk customers (inactive, declining spend)."
)

heading("3.3 Key Attributes", level=2)

add_table(
    ["Column", "Type", "Role in Dashboard"],
    [
        ["InvoiceDate", "Datetime", "Time-series aggregation by month, hour, day-of-week"],
        ["CustomerID", "String/ID", "RFM segmentation, loyalty analysis"],
        ["Country", "String", "Geographic choropleth map, top-N country ranking"],
        ["Category", "String", "Category treemap, heatmap by day of week"],
        ["Quantity", "Integer", "Order volume metrics"],
        ["UnitPrice", "Float", "Revenue computation (Quantity × UnitPrice)"],
        ["Discount", "Float", "Waterfall breakdown (discount impact)"],
        ["ShippingCost", "Float", "Profit margin calculation"],
        ["PaymentMethod", "String", "Payment method sunburst chart"],
        ["SalesChannel", "String", "Online vs. In-store channel comparison"],
        ["ReturnStatus", "String", "Return rate KPI, Sankey flow"],
    ]
)

heading("3.4 Preprocessing & Derived Columns", level=2)

body("The following transformations are applied before visualization:")

add_table(
    ["Derived Column", "Formula / Method", "Purpose"],
    [
        ["Revenue", "Quantity × UnitPrice × (1 − Discount)", "Primary financial metric for all charts"],
        ["Profit", "Revenue − ShippingCost", "Waterfall chart net value"],
        ["Month", "InvoiceDate → YYYY-MM period", "Monthly trend aggregation"],
        ["Hour", "InvoiceDate → hour (0–23)", "24-hour radial clock chart"],
        ["DayOfWeek", "InvoiceDate → Monday…Sunday", "Day × Category heatmap"],
        ["IsReturned", "ReturnStatus contains 'return'", "Return rate KPI"],
        ["IsOnline", "SalesChannel contains 'online'", "Channel split analysis"],
    ]
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 4. EDA SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
heading("4. Exploratory Data Analysis (EDA)")

body(
    "The following EDA was performed on the dataset to guide the design of the dashboard "
    "visualizations. Full code with inline charts is provided in the accompanying Jupyter "
    "notebook (milestone1.ipynb)."
)

heading("4.1 Missing Value Analysis", level=2)
body(
    "The dataset was checked for missing values across all 17 columns. Any missing values in "
    "critical fields (InvoiceDate, CustomerID, UnitPrice, Quantity) were dropped before "
    "analysis. Optional fields such as Discount and ShippingCost were filled with 0 where "
    "absent, as their absence indicates no discount/shipping was applied."
)

heading("4.2 Key Descriptive Statistics", level=2)
body(
    "Initial exploration revealed the following high-level business metrics (exact values "
    "computed from the dataset in the accompanying notebook):"
)
bullet("Total Revenue: computed as sum of Quantity × UnitPrice × (1 − Discount)")
bullet("Total Orders: total number of invoice records")
bullet("Unique Customers: number of distinct CustomerIDs")
bullet("Average Order Value: mean revenue per transaction")
bullet("Return Rate: percentage of orders with return status")
bullet("Online Channel Share: percentage of orders placed through the online channel")

body(
    "The revenue distribution was found to be right-skewed, confirming the industry observation "
    "that a small proportion of high-value orders account for a disproportionate share of total "
    "revenue. This validates the need for RFM-based customer segmentation to identify and "
    "retain those high-value customers."
)

heading("4.3 Key Patterns Discovered", level=2)

add_table(
    ["Analysis", "Finding", "Dashboard Response"],
    [
        ["Revenue by month", "Clear seasonal peaks visible in certain months", "Line/area chart with monthly granularity"],
        ["Revenue by category", "Certain categories contribute disproportionately", "Treemap + horizontal bar chart"],
        ["Top countries", "Revenue is concentrated in top 5–8 countries", "Choropleth map + top-N filter"],
        ["Hourly distribution", "Sales cluster during specific hours of the day", "24-hour polar clock chart"],
        ["Channel split", "Online vs. In-store shows meaningful revenue difference", "Sunburst + Sankey flow"],
        ["Customer frequency", "Majority are one-time buyers; small % are repeat buyers", "Loyalty bubble chart"],
        ["Revenue skewness", "Right-skewed — a few high-value orders dominate", "Box plots + RFM segmentation"],
    ]
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 5. DASHBOARD PROTOTYPE LAYOUT
# ══════════════════════════════════════════════════════════════════════════════
heading("5. Dashboard Prototype Layout")

body(
    "The prototype has been implemented in Streamlit and is accessible locally at "
    "http://localhost:8508. The layout is structured into the following sections:"
)

add_table(
    ["Section", "Charts Included", "Interaction Type"],
    [
        ["KPI Row", "6 metric cards (Revenue, Orders, Customers, AOV, Return Rate, Online %)", "Dynamic (updates with filters)"],
        ["Time Analysis", "Revenue & orders over time (line chart)", "Hover tooltips, year filter"],
        ["Category Analysis", "Treemap + grouped bar chart", "Sidebar category filter"],
        ["Geographic", "World choropleth map", "Hover by country, top-N slider"],
        ["Payment & Channel", "Sunburst chart", "Click to drill down"],
        ["Sales Journey", "Sankey flow diagram", "Hover on flow links"],
        ["Financials", "Waterfall chart (Gross → Net)", "Hover on each bar"],
        ["Time Patterns", "24-hour polar clock + day heatmap", "Hover tooltips"],
        ["Customer Loyalty", "Bubble chart (repeat buyers)", "Hover with customer stats"],
        ["Deep Exploration", "Parallel coordinates (6 dimensions)", "Brush axes to filter"],
        ["AI Assistant", "Natural language chatbot (Claude AI)", "Type any business question"],
    ]
)

body(
    "All charts respond dynamically to the sidebar filters (Country, Category, Sales Channel, "
    "Year Range), implementing the brushing and filtering interaction principles described by "
    "Heer & Shneiderman (2012)."
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 6. NEXT STEPS
# ══════════════════════════════════════════════════════════════════════════════
heading("6. Next Steps")

bullet("Add formal RFM customer segmentation with Champions / Loyal / At-Risk segment labels and a segment size visualization")
bullet("Deploy to Streamlit Cloud to generate a public shareable URL for final submission")
bullet("Add cohort retention analysis (how many customers return month over month)")
bullet("Polish narrative text panels within the dashboard for non-technical stakeholder use")
bullet("Conduct user testing with a non-technical audience to validate usability")

doc.add_paragraph()
divider()

# ══════════════════════════════════════════════════════════════════════════════
# 7. REFERENCES
# ══════════════════════════════════════════════════════════════════════════════
heading("7. References")

refs = [
    "Sarikaya, A., Correll, M., Bartram, L., Tory, M., & Fisher, D. (2019). What do we talk about when we talk about dashboards? IEEE Transactions on Visualization and Computer Graphics, 25(1), 682–692. https://doi.org/10.1109/TVCG.2018.2864903",
    "Heer, J., & Shneiderman, B. (2012). Interactive dynamics for visual analysis. Queue, 10(2), 30–55. https://doi.org/10.1145/2133416.2146416",
    "Ngai, E.W.T., Xiu, L., & Chau, D.C.K. (2009). Application of data mining techniques in customer relationship management: A literature review and classification. Expert Systems with Applications, 36(2), 2592–2602. https://doi.org/10.1016/j.eswa.2007.10.017",
    "Chen, I.J., & Popovich, K. (2003). Understanding customer relationship management (CRM): People, process and technology. Business Process Management Journal, 9(5), 672–688. https://doi.org/10.1108/14637150310496758",
    "Statista. (2024). E-commerce worldwide — Statistics & Facts. https://www.statista.com/topics/871/online-shopping/",
]

for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(f"{i}. {ref}")
    run.font.size = Pt(10)

# ── Save ──────────────────────────────────────────────────────────────────────
out = '/Users/neenubonny/online-sales-dashboard/Milestone1_OnlineSalesDashboard.docx'
doc.save(out)
print(f"Saved: {out}")
